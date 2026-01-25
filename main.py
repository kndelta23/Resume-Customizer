import functions_framework
import google.generativeai as genai
from google.cloud import storage
import pypdf
import os
import json
import tempfile
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Initialize Gemini with API Key
API_KEY = os.environ.get("GEMINI_API_KEY")
if API_KEY:
    genai.configure(api_key=API_KEY)

# Configuration
BUCKET_NAME = "knd2321-test2"
RESUME_FILENAME = "John Thomas Delta --  Resume 2025.pdf"

@functions_framework.http
def customize_resume(request):
    """HTTP Cloud Function to customize resume using Gemini, PDF from GCS, and optional Job URL."""
    
    # Handle CORS/OPTIONS
    if request.method == 'OPTIONS':
        headers = {
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'POST',
            'Access-Control-Allow-Headers': 'Content-Type',
            'Access-Control-Max-Age': '3600'
        }
        return ('', 204, headers)

    headers = {
        'Access-Control-Allow-Origin': '*'
    }

    try:
        request_json = request.get_json(silent=True)
        if not request_json:
            return (json.dumps({"error": "JSON body required"}), 400, headers)
        
        job_description = request_json.get("job_description")
        job_url = request_json.get("job_url")
        
        # Logic to fetch job description from URL if provided
        if job_url:
            try:
                print(f"Fetching job description from: {job_url}")
                page = requests.get(job_url, timeout=10)
                page.raise_for_status()
                soup = BeautifulSoup(page.content, 'html.parser')
                
                for script in soup(["script", "style", "nav", "footer", "header"]):
                    script.extract()
                
                text = soup.get_text(separator=' ')
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                cleaned_text = '\n'.join(chunk for chunk in chunks if chunk)
                
                job_description = f"Content fetched from {job_url}:\n\n{cleaned_text}"
                
            except Exception as e:
                return (json.dumps({"error": f"Failed to fetch job URL: {str(e)}"}), 400, headers)

        if not job_description:
            return (json.dumps({"error": "Either 'job_description' or 'job_url' is required."}), 400, headers)

        if not API_KEY:
             return (json.dumps({"error": "GEMINI_API_KEY environment variable not set."}), 500, headers)

        # Download PDF from GCS
        storage_client = storage.Client()
        bucket = storage_client.bucket(BUCKET_NAME)
        blob = bucket.blob(RESUME_FILENAME)
        
        with tempfile.NamedTemporaryFile(mode='wb', delete=False) as temp_pdf:
            blob.download_to_filename(temp_pdf.name)
            temp_pdf_path = temp_pdf.name

        # Extract text from PDF
        resume_text = ""
        try:
            reader = pypdf.PdfReader(temp_pdf_path)
            for page in reader.pages:
                resume_text += page.extract_text() + "\n"
        finally:
            os.remove(temp_pdf_path) # Cleanup

        if not resume_text.strip():
             return (json.dumps({"error": "Could not extract text from the resume PDF."}), 500, headers)

        # Prompt engineering
        prompt = f"""
        You are an expert career coach helping a 21-year-old recent college graduate apply for jobs.
        
        Task: Customize the following resume and write a cover letter for the provided job description.
        
        Job Description:
        {job_description}
        
        Resume:
        {resume_text}
        
        Output Requirements:
        1. Identify the Company Name and Job Title.
        2. **Relocation Logic**: Check if the job description requires the candidate to be on-site. If it does (or if it's not explicitly remote), include a sentence in the final paragraph of the cover letter expressing explicit willingness to relocate for the role.
        3. Resume Content: Rewrite the summary and experience bullet points to highlight relevant skills. 
        4. **Resume Skills**: Consolidate all skills into a single comma-separated list. Do NOT use sub-labels or categories.
        5. Cover Letter: Write a compelling cover letter in the voice of a professional, eager 21-year-old recent grad. Use the closing "Sincerely, Tommy Delta".
        6. Return the response strictly as a JSON object with this schema:
        {{
            "company": "Company Name",
            "job_title": "Job Title",
            "cover_letter_text": "Full text...",
            "resume_data": {{
                "contact_info": {{ ... }},
                "summary": "...",
                "education": [ ... ],
                "experience": [ ... ],
                "projects": [ ... ],
                "skills": "Skill 1, Skill 2, Skill 3, ..."
            }}
        }}
        """

        model = genai.GenerativeModel("gemini-flash-latest", generation_config={"response_mime_type": "application/json"})
        response = model.generate_content(prompt)
        
        try:
            response_json = json.loads(response.text)
            company = response_json.get("company", "Unknown_Company").replace(" ", "_")
            job_title = response_json.get("job_title", "Job").replace(" ", "_")
            
            # Sanitize filename components
            company = "".join([c for c in company if c.isalnum() or c in ('_', '-')])
            job_title = "".join([c for c in job_title if c.isalnum() or c in ('_', '-')])
            
            # Enforce Signature
            cl_text = response_json.get("cover_letter_text", "").strip()
            required_signature = "Sincerely,\nTommy Delta"
            # allow for slight variations like "Sincerely, \nTommy Delta"
            if "Tommy Delta" not in cl_text[-50:]: # Check end of string
                 cl_text += f"\n\n{required_signature}"
                 response_json["cover_letter_text"] = cl_text
            
        except json.JSONDecodeError:
            return (json.dumps({"error": "Failed to generate valid JSON content from model."}), 500, headers)

        # Helper to set font
        def set_font(run, size=12, bold=False, italic=False):
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic

        # --- Generate Resume DOCX ---
        doc_resume = Document()
        section = doc_resume.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        data = response_json.get("resume_data", {})
        contact = data.get("contact_info", {})

        # Header (Name & Contact)
        p = doc_resume.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact.get("name", "").upper())
        set_font(run, size=16, bold=True)
        
        p = doc_resume.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_parts = [contact.get(k) for k in ["email", "phone", "location", "linkedin"] if contact.get(k)]
        run = p.add_run(" | ".join(info_parts))
        set_font(run, size=11)
        doc_resume.add_paragraph() # Spacer

        # Helper for Section Headers
        def add_section_header(doc, title):
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            set_font(run, size=12, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run.font.underline = True

        # Summary
        if data.get("summary"):
            add_section_header(doc_resume, "Professional Summary")
            p = doc_resume.add_paragraph(data["summary"])
            set_font(p.runs[0], size=12)

        # Education
        if data.get("education"):
            add_section_header(doc_resume, "Education")
            for edu in data["education"]:
                # Line 1: Institution -- Location
                p = doc_resume.add_paragraph()
                run_inst = p.add_run(edu.get("institution", ""))
                set_font(run_inst, size=12, bold=True)
                if edu.get("location"):
                    run_loc = p.add_run(f" \t{edu['location']}")
                    set_font(run_loc, size=12)
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                
                # Line 2: Degree -- Date
                p = doc_resume.add_paragraph()
                run_deg = p.add_run(edu.get("degree", ""))
                set_font(run_deg, size=12, italic=True)
                if edu.get("graduation_date"):
                    run_date = p.add_run(f" -- {edu['graduation_date']}")
                    set_font(run_date, size=12)
                
                # Details
                for detail in edu.get("details", []):
                     p = doc_resume.add_paragraph(detail, style='List Bullet')
                     set_font(p.runs[0], size=12)

        # Experience
        if data.get("experience"):
            add_section_header(doc_resume, "Experience")
            for exp in data["experience"]:
                # Line 1: Company -- Location
                p = doc_resume.add_paragraph()
                run_comp = p.add_run(exp.get("company", ""))
                set_font(run_comp, size=12, bold=True)
                
                # Line 2: Title -- Dates
                p = doc_resume.add_paragraph()
                run_title = p.add_run(exp.get("title", ""))
                set_font(run_title, size=12, italic=True)
                if exp.get("dates"):
                    run_dates = p.add_run(f" ({exp['dates']})")
                    set_font(run_dates, size=12)

                # Bullets
                for bullet in exp.get("bullets", []):
                    p = doc_resume.add_paragraph(bullet, style='List Bullet')
                    if p.runs: set_font(p.runs[0], size=12)

        # Projects
        if data.get("projects"):
            add_section_header(doc_resume, "Projects")
            for proj in data["projects"]:
                p = doc_resume.add_paragraph()
                run_name = p.add_run(proj.get("name", ""))
                set_font(run_name, size=12, bold=True)
                if proj.get("dates"):
                     run_date = p.add_run(f" | {proj['dates']}")
                     set_font(run_date, size=12)
                
                for bullet in proj.get("bullets", []):
                    p = doc_resume.add_paragraph(bullet, style='List Bullet')
                    if p.runs: set_font(p.runs[0], size=12)

        # Skills (Consolidated)
        skills_text = data.get("skills")
        if skills_text:
            add_section_header(doc_resume, "Skills")
            # Handle if model returns dict by mistake, or string
            if isinstance(skills_text, dict):
                 # Fallback: flatten values
                 skills_str = ", ".join([v for v in skills_text.values()])
            elif isinstance(skills_text, list):
                 skills_str = ", ".join(skills_text)
            else:
                 skills_str = str(skills_text)

            p = doc_resume.add_paragraph(skills_str)
            set_font(p.runs[0], size=12)

        # --- Generate Cover Letter DOCX ---
        doc_cl = Document()
        section = doc_cl.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        p = doc_cl.add_paragraph(response_json.get("cover_letter_text", ""))
        set_font(p.runs[0], size=12)
        
        # --- Save and Upload ---
        def save_and_upload(doc, file_prefix):
            filename = f"{company}_{job_title}_{file_prefix}_{datetime.now().strftime('%Y-%m-%d')}.docx"
            temp_path = f"/tmp/{filename}"
            doc.save(temp_path)
            
            try:
                output_blob = bucket.blob(filename)
                output_blob.upload_from_filename(temp_path)
                return f"gs://{BUCKET_NAME}/{filename}", filename
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

        resume_uri, resume_filename = save_and_upload(doc_resume, "Resume")
        cl_uri, cl_filename = save_and_upload(doc_cl, "CoverLetter")

        return (json.dumps({
            "result": "Success",
            "resume_docx": resume_uri,
            "resume_filename": resume_filename,
            "cover_letter_docx": cl_uri,
            "cover_letter_filename": cl_filename
        }), 200, headers)

    except Exception as e:
        return (json.dumps({"error": str(e)}), 500, headers)
